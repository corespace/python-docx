# encoding: utf-8

"""
Step implementations for header-related features
"""

from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

from behave import given, then, when

from docx import Document
from docx.enum.header import WD_HEADER_FOOTER
from helpers import test_docx, saved_docx_path, test_text


# given ===================================================

@given('a header and footer {having_or_no} definition')
def given_a_header_having_or_no_definition(context, having_or_no):
    filename = {
        'having a':  'hdr-header-props',
        'having no': 'doc-default',
    }[having_or_no]
    document = Document(test_docx(filename))
    context.document = document
    context.header = document.sections[0].headers[WD_HEADER_FOOTER.PRIMARY]
    context.footer = document.sections[0].footers[WD_HEADER_FOOTER.PRIMARY]


# when =====================================================

@when('I change the text in the footer')
def i_change_the_text_in_the_footer(context):
    context.footer.body.paragraphs[0].text = test_text


# then =====================================================

@then('header.body contains the text of the header')
def then_header_body_contains_the_text_of_the_header(context):
    header = context.header
    text = header.body.paragraphs[0].text
    assert text == 'S1HP1'


@then('footer.body contains the text of the footer')
def then_footer_body_contains_the_text_of_the_header(context):
    text = context.footer.body.paragraphs[0].text
    print(text)
    assert text == 'S1FP1'


@then('header.body is a HeaderFooterBody object')
def then_header_body_is_a_HeaderFooterBody_object(context):
    header = context.header
    assert type(header.body).__name__ == 'HeaderFooterBody'


@then('footer.body is a HeaderFooterBody object')
def then_footer_body_is_a_HeaderFooterBody_object(context):
    assert type(context.footer.body).__name__ == 'HeaderFooterBody'


@then('header.is_linked_to_previous is {value}')
def then_header_is_linked_to_previous_is_value(context, value):
    expected_value = {'True': True, 'False': False}[value]
    header = context.header
    assert header.is_linked_to_previous is expected_value


@then('footer.is_linked_to_previous is {value}')
def then_footer_is_linked_to_previous_is_value(context, value):
    expected_value = {'True': True, 'False': False}[value]
    assert context.footer.is_linked_to_previous is expected_value


@then('the footer contains the text I added')
def the_footer_contains_the_text_i_added(context):
    document = Document(saved_docx_path)
    footer = document.sections[0].footers[WD_HEADER_FOOTER.PRIMARY]
    assert footer.body.paragraphs[0].text == test_text
